# generate_docs.py
"""Генерация README.md и .docx документации из исходников проекта.

Универсальный инструмент:
- Как standalone: python generate_docs.py  (редактируй секцию НАСТРОЙКИ)
- Как библиотека: from generate_docs import generate_docs_markdown
"""

import os
import re
from pathlib import Path
from datetime import datetime

import pathspec

try:
    from docx import Document
    from docx.shared import Pt
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

# ══════════════════════════════════════════════════════════════════════════════
# НАСТРОЙКИ ПРОЕКТА  (редактируй при миграции, используются только в standalone)
# ══════════════════════════════════════════════════════════════════════════════

PROJECT_ROOT = Path(__file__).parent

# Краткое описание проекта — подставляется в README
DESCRIPTION = "Автоматизация полного цикла разработки на нескольких языках программирования."

# ══════════════════════════════════════════════════════════════════════════════

# Расширения файлов для включения в документацию
INCLUDED_EXTENSIONS = [
    ".py", ".yaml", ".yml", ".toml", ".json",
    ".sh", ".sql", ".txt", ".md", ".env",
    ".js", ".ts", ".jsx", ".tsx", ".css", ".html",
    ".go", ".rs", ".rb", ".java", ".kt", ".c", ".cpp", ".h",
]

# Файлы без расширения, которые нужно включить
EXTRA_FILES_WITHOUT_EXTENSION = {
    "Dockerfile", "Dockerfile.cpu", "Dockerfile.gpu",
    "Makefile", ".dockerignore", "LICENSE",
}

# Файлы, которые всегда исключаются (секреты, локальные конфиги)
EXCLUDED_FILES = {
    ".env", "secrets.py", "config_local.py",
    "id_rsa", "id_rsa.pub", "known_hosts",
    "docker-compose.override.yml", "token.txt",
}

SYNTAX_MAP = {
    ".py": "python",
    ".yaml": "yaml",
    ".yml": "yaml",
    ".sh": "bash",
    ".env": "env",
    ".txt": "",
    ".md": "markdown",
    ".json": "json",
    ".toml": "toml",
    ".sql": "sql",
    ".js": "javascript",
    ".ts": "typescript",
    ".jsx": "jsx",
    ".tsx": "tsx",
    ".css": "css",
    ".html": "html",
    ".go": "go",
    ".rs": "rust",
    ".rb": "ruby",
    ".java": "java",
    ".kt": "kotlin",
    ".c": "c",
    ".cpp": "cpp",
    ".h": "c",
    "Dockerfile": "dockerfile",
    "Makefile": "makefile",
}


# ── Вспомогательные функции ──────────────────────────────────────────────────

def _load_gitignore_spec(project_root: Path):
    gitignore_path = project_root / ".gitignore"
    if not gitignore_path.exists():
        return None
    with open(gitignore_path, "r", encoding="utf-8") as f:
        return pathspec.PathSpec.from_lines("gitwildmatch", f)


def _should_include(path: Path, project_root: Path, spec) -> bool:
    if path.name in EXCLUDED_FILES:
        return False
    rel = path.relative_to(project_root)
    return spec is None or not spec.match_file(rel)


def _is_doc_file(name: str) -> bool:
    if name in EXTRA_FILES_WITHOUT_EXTENSION:
        return True
    return any(name.endswith(ext) for ext in INCLUDED_EXTENSIONS)


def _generate_tree(start_path: Path, project_root: Path, spec) -> str:
    lines = []
    for root, dirs, files in os.walk(start_path):
        dirs[:] = sorted(d for d in dirs if _should_include(Path(root) / d, project_root, spec))
        level = len(Path(root).relative_to(start_path).parts)
        indent = "    " * level
        lines.append(f"{indent}{Path(root).name}/")
        sub_indent = "    " * (level + 1)
        for f in sorted(files):
            if _should_include(Path(root) / f, project_root, spec) and _is_doc_file(f):
                lines.append(f"{sub_indent}{f}")
    return "\n".join(lines)


def _read_file_content(file_path: Path) -> str:
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        if file_path.name in SYNTAX_MAP:
            lang = SYNTAX_MAP[file_path.name]
        else:
            lang = SYNTAX_MAP.get(file_path.suffix, "")
        return f"```{lang}\n{content}\n```"
    except Exception as e:
        return f"<!-- Ошибка чтения файла: {e} -->\n\n"


def _gather_files(start_path: Path, project_root: Path, spec) -> str:
    sections = []
    for root, dirs, files in os.walk(start_path):
        dirs[:] = sorted(d for d in dirs if _should_include(Path(root) / d, project_root, spec))
        for f in sorted(files):
            if f == "README.md":
                continue
            full_path = Path(root) / f
            if _should_include(full_path, project_root, spec) and _is_doc_file(f):
                rel = full_path.relative_to(project_root)
                sections.append(f"### `{rel}`\n")
                sections.append(_read_file_content(full_path))
    return "\n".join(sections)


def _build_header(name: str, description: str, structure: str) -> str:
    return f"""# {name}

## Описание

{description}

## Установка

1. Клонируйте репозиторий
2. Создайте виртуальное окружение: `python -m venv .venv`
3. Активируйте:
   - Windows: `.venv\\Scripts\\activate`
   - Linux/macOS: `source .venv/bin/activate`
4. Установите зависимости: `pip install -r requirements.txt`

## Запуск

Описан в файле RUN.md

## Структура проекта

{structure}
"""


# ── Публичное API ────────────────────────────────────────────────────────────

def generate_docs_markdown(
    project_path: Path, description: str = "", name: str = "",
) -> str:
    """Генерирует markdown-документацию для проекта. Возвращает строку."""
    spec = _load_gitignore_spec(project_path)
    structure = _generate_tree(project_path, project_path, spec)
    code_sections = _gather_files(project_path, project_path, spec)

    md = _build_header(name or project_path.name, description, structure)
    md += f"## Коды основных модулей\n\n{code_sections}\n"
    return md


# ── Docx ─────────────────────────────────────────────────────────────────────

def _clean_for_xml(text: str) -> str:
    if not isinstance(text, str):
        text = str(text)
    text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text)
    text = re.sub(r"[\uFFFE\uFFFF]", "", text)
    try:
        text = text.encode("utf-8", errors="ignore").decode("utf-8", errors="ignore")
    except Exception:
        text = "".join(ch for ch in text if ord(ch) >= 32)
    return text


def _write_docx(documentation: str, project_root: Path) -> None:
    if not _HAS_DOCX:
        print("[~] python-docx не установлен — .docx не генерируется")
        return

    docs_dir = project_root / "documents"
    docs_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y_%m_%d_%H_%M")
    docx_path = docs_dir / f"{project_root.name}_{timestamp}.docx"

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font  # type: ignore[union-attr]
    font.name = "Consolas"
    font.size = Pt(10)

    lines = documentation.splitlines()
    i = 0
    while i < len(lines):
        line = _clean_for_xml(lines[i])
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(_clean_for_xml(lines[i]))
                i += 1
            if code_lines:
                doc.add_paragraph("\n".join(code_lines))
            i += 1
            continue
        if line:
            p = doc.add_paragraph(line)
            p.style = doc.styles["Normal"]  # type: ignore[assignment]
        i += 1

    try:
        doc.save(docx_path)  # type: ignore[arg-type]
        print(f"[+] Документация сохранена: {docx_path}")
    except Exception as e:
        print(f"[!] Ошибка при сохранении .docx: {e}")


# ── Standalone ───────────────────────────────────────────────────────────────

def build_documentation():
    """Standalone: генерирует README.md + .docx для PROJECT_ROOT."""
    md = generate_docs_markdown(PROJECT_ROOT, DESCRIPTION)

    readme_file = PROJECT_ROOT / "README.md"
    with open(readme_file, "w", encoding="utf-8") as f:
        f.write(md)
    print("[+] README.md обновлён")

    _write_docx(md, PROJECT_ROOT)


if __name__ == "__main__":
    build_documentation()
    print("[+] Генерация документации завершена")
